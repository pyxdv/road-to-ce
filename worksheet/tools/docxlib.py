# -*- coding: utf-8 -*-
"""ตัวช่วยสร้างใบงาน .docx ให้ตรงกับเทมเพลตเดิมของรายวิชา (TH Sarabun New 16pt)"""
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Twips

FONT = "TH Sarabun New"
SZ = 32          # half-points -> 16pt
SZ_SMALL = 28    # 14pt


def _rpr(run, size, bold, color, italic=False):
    rpr = run._r.get_or_add_rPr()
    fonts = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        fonts.set(qn(a), FONT)
    rpr.append(fonts)
    if bold:
        for tag in ("w:b", "w:bCs"):
            rpr.append(OxmlElement(tag))
    if italic:
        for tag in ("w:i", "w:iCs"):
            rpr.append(OxmlElement(tag))
    for tag in ("w:sz", "w:szCs"):
        e = OxmlElement(tag)
        e.set(qn("w:val"), str(size))
        rpr.append(e)
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rpr.append(c)
    rpr.append(OxmlElement("w:cs"))   # บอก Word ว่าข้อความนี้เป็น complex script (ไทย)


def para(container, text="", bold=False, size=SZ, color=None, indent=0,
         align=None, before=0, after=0, italic=False, line=None):
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line
    if indent:
        pf.left_indent = Twips(indent)
    if align == "center":
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if text:
        run = p.add_run(text)
        _rpr(run, size, bold, color, italic)
    return p


def rich(container, parts, size=SZ, indent=0, before=0, after=0, align=None):
    """parts = [(text, bold, color), ...] รวมหลายรูปแบบในย่อหน้าเดียว"""
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent:
        pf.left_indent = Twips(indent)
    if align == "center":
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for item in parts:
        text, bold = item[0], item[1]
        color = item[2] if len(item) > 2 else None
        run = p.add_run(text)
        _rpr(run, size, bold, color)
    return p


def bullets(container, items, marker="•", indent=340, size=SZ, after=0):
    for it in items:
        if isinstance(it, tuple):
            head, tail = it
            rich(container, [(f"{marker} {head}", True), (tail, False)],
                 size=size, indent=indent, after=after)
        else:
            para(container, f"{marker} {it}", size=size, indent=indent, after=after)


def numbered(container, items, indent=340, size=SZ, after=0, start=1):
    for i, it in enumerate(items, start):
        if isinstance(it, tuple):
            head, tail = it
            rich(container, [(f"{i}. {head}", True), (tail, False)],
                 size=size, indent=indent, after=after)
        else:
            para(container, f"{i}. {it}", size=size, indent=indent, after=after)


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def cell_text(cell, text, bold=False, size=SZ, color=None, align=None, first=True):
    p = cell.paragraphs[0] if first and not cell.paragraphs[0].text else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if align == "center":
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        _rpr(p.add_run(text), size, bold, color)
    return p


def grid(container, rows, cols, widths=None):
    t = container.add_table(rows=rows, cols=cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if widths:
        t.autofit = False
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Twips(w)
    return t


def box(container, width=9154):
    """กล่องเส้นเดียว 1x1 แบบที่ใบงานเดิมใช้ครอบเนื้อหา"""
    t = grid(container, 1, 1, [width])
    cell = t.rows[0].cells[0]
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    return cell


def blank_lines(container, n, indent=0, size=SZ, after=6):
    for _ in range(n):
        para(container, "_" * 92, size=size, indent=indent, after=after)


def page_break(container):
    p = container.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    return p


def clear_body(doc):
    body = doc.element.body
    sect = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not sect:
            body.remove(child)


def set_header(doc, replacements):
    """แก้ข้อความในหัวกระดาษ เช่น เลขภาคเรียน หรือชื่อชนิดเอกสาร

    replacements = {"1/2568": "1/2569", "ใบงานที่ ": "ใบความรู้ที่ "}
    """
    hdr = doc.sections[0].header
    # เซลล์ที่ถูก merge จะถูกวนซ้ำหลายรอบ ถ้าไม่กันไว้ ข้อความจะถูกแทนที่ซ้ำ
    # เช่น "ใบงานที่" -> "เฉลยใบงานที่" -> "เฉลยเฉลยใบงานที่"
    seen = set()
    for para_or_cell in _iter_header_paragraphs(hdr):
        for run in para_or_cell.runs:
            if id(run._r) in seen:
                continue
            seen.add(id(run._r))
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)


def _iter_header_paragraphs(hdr):
    for p in hdr.paragraphs:
        yield p
    for table in hdr.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
